import re

import requests
from docx.api import Document

from locators import PurchaseSupplierResultsLocators, DocumentsResultsLocators
from logger_settings import logger
from pages.base_page import BasePage


class DocumentsResults(BasePage):
    def __init__(self, purchase_search_page_tree, order_num):
        self.order_num = order_num
        self.purchase_search_page_tree = purchase_search_page_tree

    def get_status(self):
        if not self.check_element_existing(PurchaseSupplierResultsLocators.status, self.purchase_search_page_tree):
            return ''
        else:

            status = self.purchase_search_page_tree.xpath(PurchaseSupplierResultsLocators.status)[0].lstrip().rstrip()
        return status

    def download_doc(self, url):
        headers = {
            'User-Agent': 'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/102.0.0.0 Safari/537.36',
        }
        response = requests.get(url, allow_redirects=True, headers=headers)
        fname = re.findall("filename=(.+)", response.headers['content-disposition'])[0].lower()
        if '.docx' in fname:
            fname = "doc.docx"
            # TODO: статичный путь сделать
            with open(fname, "wb") as file:
                file.write(response.content)
                return fname

        elif '.doc' in fname:
            logger.info('".doc" format. Skip oppening')
        elif '.rar' in fname or '.zip' in fname:
            logger.debug('Archive. Skipping')
        else:
            logger.critical(f"Неизвестный формат файла:{fname}")

        return None


    def get_contract_file(self, doc_block_tree):
        if not self.check_element_existing(DocumentsResultsLocators.a_contract_file, doc_block_tree):
            return ''
        else:
            status = doc_block_tree.xpath(DocumentsResultsLocators.a_contract_file)[0].lstrip().rstrip()
        return status

    def get_ru_from_table(self, document):
        table_with_ru = self.find_table_with_ru(document.tables)

        if table_with_ru:
            ru = []

            for row in table_with_ru.rows:
                for cell in row.cells:
                    text = cell.text
                    if re.search(r'[2]0[0-2][0-9]/', text):
                        year = re.findall(r'[2]0[0-2][0-9]', text)[0]
                        remainder = text.split(f'{year}/')[1].split()[0]
                        ru.append(f'{year}/{remainder}')

            uniq_ru = list(set(ru))
            return uniq_ru

        else:
            return []

    def get_registry_entry_numbers_from_table(self, document):

        table_with_registry_entry_numbers = self.find_table_with_registry_entry_numbers(
            document.tables)

        if table_with_registry_entry_numbers:
            registry_entry_numbers = []

            for row in table_with_registry_entry_numbers.rows:
                for cell in row.cells:
                    text = rf'{cell.text}'

                    if re.search(r'\d*\\\d*\\[2]0[0-2][0-9]', text):
                        registry_entry_numbers.append(re.findall(r'\d*\\\d*\\[2]0[0-2][0-9]', text)[0])

            uniq_registry_entry_numbers = list(set(registry_entry_numbers))
            return uniq_registry_entry_numbers

        else:
            return []

    def find_table_with_ru(self, tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text

                    if re.search(r'[2]0[0-2][0-9]/', text):
                        return table

        return ''

    def find_table_with_registry_entry_numbers(self, tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text

                    if re.search(r'\d*\\\d*\\[2]0[0-2][0-9]', text):
                        return table

        return ''

    def get_draft_id(self):
        if not self.check_element_existing(PurchaseSupplierResultsLocators.data_id_draft_id,
                                           self.purchase_search_page_tree):
            return ''
        else:
            draft_id = self.purchase_search_page_tree.xpath(PurchaseSupplierResultsLocators.data_id_draft_id)[0]
        return draft_id

    def get_doc_block_tree(self):
        draft_id = self.get_draft_id()
        if draft_id:

            tree = self.get_tree(
                # разворачиваем часть страницы "Информация о процедуре заключения контракта"
                'https://zakupki.gov.ru/epz/order/notice/rpec/documents-results.html',
                {'orderNum': self.order_num,
                 'draftId': draft_id  # idшник документа, судя во всему. Без него возвращает 404
                 }
            )
            return tree
        else:
            return ''

    def get_ru_and_registry_entry_numbers(self):
        # TODO: возможно в будущем окажется, что может быть несколько документов. Нужно будет перепилить
        status = self.get_status()

        doc_block_tree = self.get_doc_block_tree()
        if doc_block_tree != '':
            if status == 'Контракт не заключен':
                self.ru_number = []
            else:
                doc_link = self.get_contract_file(doc_block_tree)
                if doc_link:
                    doc_name = self.download_doc(doc_link)
                    if doc_name:
                        try:
                            document = Document(doc_name)

                            ru = self.get_ru_from_table(document)
                            registry_entry_numbers = self.get_registry_entry_numbers_from_table(document)

                            return ru, registry_entry_numbers

                        # пример ошибки:
                        # ValueError: file 'doc.docx' is not a Word file, content type is 'application/vnd.openxmlformats-officedocument.themeManager+xml'
                        # закупка, ее вызвавшая: https://zakupki.gov.ru/epz/order/notice/ea20/view/supplier-results.html?regNumber=0351100025322000002
                        except ValueError:
                            return [], []

        return [], []
